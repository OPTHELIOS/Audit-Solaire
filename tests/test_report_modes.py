"""Tests pour la divergence audit_complet vs diagnostic_court."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from domain.audit_studio import ModeRapport
from domain.docx_service import build_docx_report
from domain.energy import EnergyInputs, compute_energy
from domain.models import Audit, Preuve, TypePreuve
from domain.report_service import build_report_markdown


def _build_audit(mode: ModeRapport) -> Audit:
    audit = Audit()
    audit.meta.numero_audit = "AUD-TEST-MODE"
    audit.projet.operation = "Site test"
    audit.projet.maitre_ouvrage = "MOA Test"
    audit.set_mode_rapport(mode)
    audit.studio.upsert_scenario("redimensionner", True, "Surchauffes estivales.")
    audit.studio.note_strategique = "Cap métrologie avant tout investissement."

    audit.energy.inputs = EnergyInputs(
        volume_ecs_jour_litres=2000,
        delta_t_kelvin=40,
        surface_capteurs_m2=30,
        volume_stockage_solaire_litres=2000,
        zone_climatique="H2",
    )
    audit.energy.results = compute_energy(audit.energy.inputs)
    return audit


class TestMarkdownDivergence(unittest.TestCase):
    def test_audit_complet_title(self) -> None:
        audit = _build_audit(ModeRapport.audit_complet)
        md = build_report_markdown({"audit": audit})
        self.assertIn("Rapport d'audit technique solaire thermique", md)

    def test_diagnostic_court_title(self) -> None:
        audit = _build_audit(ModeRapport.diagnostic_court)
        md = build_report_markdown({"audit": audit})
        self.assertIn("Diagnostic court", md)
        self.assertIn("Tableau priorisé et chiffré", md)
        self.assertIn("Conclusion opérationnelle", md)

    def test_both_contain_energy_block(self) -> None:
        for mode in (ModeRapport.audit_complet, ModeRapport.diagnostic_court):
            audit = _build_audit(mode)
            md = build_report_markdown({"audit": audit})
            self.assertIn("Calculs énergétiques", md, msg=f"missing in {mode.value}")
            self.assertIn("Taux de couverture", md, msg=f"missing in {mode.value}")

    def test_diagnostic_excludes_detailed_findings_section(self) -> None:
        audit = _build_audit(ModeRapport.diagnostic_court)
        md = build_report_markdown({"audit": audit})
        # The detailed "Constats par section" header is exclusive to audit complet.
        self.assertNotIn("## Constats par section", md)


class TestDocxModes(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def _texts(self, path: Path) -> str:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    def test_docx_audit_complet_contains_long_sections(self) -> None:
        audit = _build_audit(ModeRapport.audit_complet)
        output = Path(self.tmp.name) / "audit.docx"
        build_docx_report({"audit": audit}, output)
        texts = self._texts(output)
        # Sections exclusives à l'audit complet (numérotées).
        self.assertIn("0. Description de l'installation", texts)
        self.assertIn("3. Lecture par section", texts)
        self.assertIn("4. Constats détaillés", texts)

    def test_docx_diagnostic_court_is_resserre(self) -> None:
        audit = _build_audit(ModeRapport.diagnostic_court)
        output = Path(self.tmp.name) / "diag.docx"
        build_docx_report({"audit": audit}, output)
        texts = self._texts(output)
        self.assertIn("Tableau de priorisation", texts)
        self.assertIn("Conclusion opérationnelle", texts)
        self.assertNotIn("4. Constats détaillés", texts)
        self.assertNotIn("0. Description de l'installation", texts)

    def test_docx_cover_page_branding(self) -> None:
        audit = _build_audit(ModeRapport.audit_complet)
        output = Path(self.tmp.name) / "cover.docx"
        build_docx_report({"audit": audit}, output)
        texts = self._texts(output)
        self.assertIn("OPT'HELIOS", texts)
        self.assertIn("AUDIT SOLAIRE THERMIQUE", texts)
        self.assertIn("contact@opthelios.fr", texts)

    def test_docx_includes_energy_section(self) -> None:
        audit = _build_audit(ModeRapport.audit_complet)
        output = Path(self.tmp.name) / "energy.docx"
        build_docx_report({"audit": audit}, output)
        texts = self._texts(output)
        self.assertIn("Calculs énergétiques", texts)
        self.assertIn("Énergie ECS annuelle", texts)


class TestPhotoCaptions(unittest.TestCase):
    def test_caption_appears_in_diagnostic_markdown(self) -> None:
        audit = _build_audit(ModeRapport.diagnostic_court)
        audit.preuves.append(
            Preuve(
                type_preuve=TypePreuve.photo,
                nom_fichier="champ.jpg",
                chemin_fichier="/tmp/champ.jpg",
                legende="Vue du champ capteurs côté sud",
            )
        )
        md = build_report_markdown({"audit": audit})
        self.assertIn("Photos clés", md)
        self.assertIn("Vue du champ capteurs côté sud", md)


if __name__ == "__main__":
    unittest.main()
