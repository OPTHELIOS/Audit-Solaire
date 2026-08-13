from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from domain.audit_studio import (
    FORMULATIONS_BY_CODE,
    MODE_RAPPORT_LABELS,
    SCENARIOS_BY_CODE,
    AuditStudioBlock,
    extract_studio_from_session,
)
from domain.control_service import group_controls_by_section
from domain.report_service import build_report_data


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).strip().split())


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_document_language(document: Document, lang_code: str = "fr-FR") -> None:
    styles = document.styles
    for style in styles:
        try:
            rpr = style.element.get_or_add_rPr()
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            lang.set(qn("w:val"), lang_code)
            lang.set(qn("w:eastAsia"), lang_code)
            lang.set(qn("w:bidi"), lang_code)
        except Exception:
            continue


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def _set_default_font(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 15, True),
        ("Heading 2", 12, True),
        ("Heading 3", 10.5, True),
    ]:
        if style_name in styles:
            styles[style_name].font.name = "Calibri"
            styles[style_name].font.size = Pt(size)
            styles[style_name].font.bold = bold


def _add_header(document: Document, report_title: str, reference: str | None = None) -> None:
    section = document.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text = report_title
    if reference:
        text += f" — {reference}"
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8)


def _add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rapport généré automatiquement")
    run.font.size = Pt(8)


LOGO_PATH = "assets/opthelios_logo.png"


def _add_cover_page(
    document: Document,
    report_title: str,
    site_name: str | None = None,
    reference: str | None = None,
    audit_date: str | None = None,
    cover_photo_path: str | None = None,
) -> None:
    """Vraie page de garde (logo, titre, photo du site si fournie, infos
    clés), au lieu d'un simple bloc de titre en haut de la première page de
    contenu. Suivie d'un saut de page (voir `build_docx_report`) : le
    rapport commence vraiment à "1. Appréciation globale" sur une page
    fraîche."""
    if Path(LOGO_PATH).exists():
        _add_picture_if_exists(document, LOGO_PATH, width_inches=1.8)
        document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report_title)
    run.bold = True
    run.font.size = Pt(22)

    if site_name:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(site_name)
        run.bold = True
        run.font.size = Pt(14)

    if cover_photo_path:
        document.add_paragraph("")
        _add_picture_if_exists(document, cover_photo_path, width_inches=4.5)

    document.add_paragraph("")

    meta_lines = []
    if reference:
        meta_lines.append(f"Référence : {reference}")
    if audit_date:
        meta_lines.append(f"Date : {audit_date}")
    else:
        meta_lines.append(f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    for line in meta_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    document.add_paragraph("")


def _add_signature_section(document: Document, auditeur: str | None = None) -> None:
    """Bloc de validation en fin de rapport : identifie l'auditeur et laisse
    une ligne pour une signature manuscrite (le document reste un .docx
    éditable, pas de signature électronique intégrée)."""
    document.add_heading("Validation", level=1)
    document.add_paragraph(
        "Ce rapport a été établi par OPT'HELIOS sur la base des constats réalisés lors "
        "de la visite d'audit et des éléments transmis par le client."
    )

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    _set_cell_text(table.rows[0].cells[0], "Auditeur", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Date", bold=True)
    _shade_cell(table.rows[0].cells[0], "D9EAF7")
    _shade_cell(table.rows[0].cells[1], "D9EAF7")

    _set_cell_text(table.rows[1].cells[0], _safe_str(auditeur))
    _set_cell_text(table.rows[1].cells[1], datetime.now().strftime("%d/%m/%Y"))

    document.add_paragraph("")
    p = document.add_paragraph()
    p.add_run("Signature :").bold = True
    document.add_paragraph("")
    document.add_paragraph("_" * 40)


def _add_bullets(document: Document, lines: list[str]) -> None:
    for line in lines:
        document.add_paragraph(line, style="List Bullet")


def _resolve_expert_conclusion(payload: dict[str, Any]) -> str:
    return _safe_str(
        payload.get("expert_conclusion")
        or payload.get("global_assessment", {}).get("commentaire_global", "")
    )


def _add_global_assessment(document: Document, payload: dict[str, Any]) -> None:
    ga = payload["global_assessment"]
    counts = payload["counts"]

    document.add_heading("1. Appréciation globale", level=1)
    document.add_paragraph(_safe_str(ga["commentaire_global"]))

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    headers = [
        "Statut global",
        "Taux de complétion",
        "Taux de conformité",
        "Constats critiques",
        "Constats majeurs",
    ]
    for cell, label in zip(hdr, headers):
        _set_cell_text(cell, label, bold=True)
        _shade_cell(cell, "D9EAF7")

    row = table.add_row().cells
    values = [
        ga["statut_global"],
        f"{ga['taux_completion_pct']} %",
        f"{ga['taux_conformite_pct']} %",
        str(counts["critical_findings"]),
        str(counts["major_findings"]),
    ]
    for cell, value in zip(row, values):
        _set_cell_text(cell, value)


def _add_executive_summary(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("2. Synthèse exécutive", level=1)
    _add_bullets(document, payload["executive_summary"])

    document.add_heading("2.1 Messages clés", level=2)
    _add_bullets(document, payload["key_messages"])

    document.add_heading("2.2 Note méthodologique", level=2)
    _add_bullets(document, payload["methodology_note"])


def _add_expert_conclusion_section(document: Document, payload: dict[str, Any]) -> None:
    expert_conclusion = _resolve_expert_conclusion(payload)
    if not expert_conclusion:
        return

    document.add_heading("2.3 Conclusion experte", level=2)
    document.add_paragraph(expert_conclusion)


def _add_section_summary(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("3. Lecture par section", level=1)

    rows = payload["section_summaries"]
    if not rows:
        document.add_paragraph("Aucun constat structurant n’est disponible à ce stade.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Section",
        "Constats",
        "Critiques",
        "Majeures",
        "Non conformes",
        "Non vérifiables",
    ]
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, label, bold=True)
        _shade_cell(cell, "D9EAF7")

    for item in rows:
        row = table.add_row().cells
        values = [
            item["section"],
            str(item["nb_constats"]),
            str(item["nb_critiques"]),
            str(item["nb_majeures"]),
            str(item["nb_non_conformes"]),
            str(item["nb_non_verifiables"]),
        ]
        for cell, value in zip(row, values):
            _set_cell_text(cell, value)

    document.add_paragraph("")

    for item in rows:
        document.add_heading(item["section"], level=2)
        document.add_paragraph(item["texte_intro"])


def _is_image_file(path: str | Path) -> bool:
    suffix = Path(path).suffix.lower()
    return suffix in {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tif", ".tiff", ".webp"}


def _add_picture_if_exists(document: Document, image_path: str | Path, width_inches: float = 2.2) -> bool:
    p = Path(image_path)
    if not p.exists():
        return False
    if not _is_image_file(p):
        return False

    try:
        document.add_picture(str(p), width=Inches(width_inches))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        return False


def _add_findings(document: Document, payload: dict[str, Any], include_evidences: bool = True) -> None:
    document.add_heading("4. Constats détaillés", level=1)

    findings_by_section = payload["findings_by_section"]
    if not findings_by_section:
        document.add_paragraph("Aucun constat détaillé n’est disponible.")
        return

    for section, rows in findings_by_section.items():
        document.add_heading(section, level=2)

        for row in rows:
            document.add_heading(f"{row['controle_id']} — {row['libelle']}", level=3)

            p = document.add_paragraph()
            r = p.add_run("Verdict : ")
            r.bold = True
            p.add_run(_safe_str(row["verdict"]))

            p = document.add_paragraph()
            r = p.add_run("Criticité : ")
            r.bold = True
            p.add_run(_safe_str(row["criticite"]))

            document.add_paragraph(_safe_str(row["phrase_constat"]))
            document.add_paragraph(_safe_str(row["phrase_impact"]))
            document.add_paragraph(_safe_str(row["phrase_action"]))

            if row.get("preuve_documentaire"):
                p = document.add_paragraph()
                r = p.add_run("Preuve documentaire : ")
                r.bold = True
                p.add_run(_safe_str(row["preuve_documentaire"]))

            if include_evidences and row.get("photos"):
                document.add_paragraph("Preuves photographiques / pièces jointes disponibles :")
                added_any = False
                for photo_path in row["photos"][:4]:
                    added = _add_picture_if_exists(document, photo_path, width_inches=2.2)
                    added_any = added_any or added
                    if not added:
                        document.add_paragraph(f"- {photo_path}", style="List Bullet")

                if not added_any:
                    document.add_paragraph(
                        "Les fichiers associés n’ont pas pu être intégrés comme images dans le document."
                    )


def _add_action_plan(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("5. Plan d’actions", level=1)

    rows = payload["action_plan"]
    if not rows:
        document.add_paragraph("Aucune action corrective n’est actuellement générée.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Priorité", "ID", "Section", "Objet", "Impact", "Action recommandée"]
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, label, bold=True)
        _shade_cell(cell, "D9EAF7")

    for item in rows:
        row = table.add_row().cells
        values = [
            item["priorite"],
            item["controle_id"],
            item["section"],
            item["objet"],
            item["impact"],
            item["action_recommandee"],
        ]
        for cell, value in zip(row, values):
            _set_cell_text(cell, _safe_str(value))


def _add_studio_sections(document: Document, studio: AuditStudioBlock | None) -> None:
    if studio is None:
        return

    document.add_heading("6. Studio OPT'HELIOS — orientations stratégiques", level=1)

    mode_label = MODE_RAPPORT_LABELS.get(studio.mode_rapport.value, studio.mode_rapport.value)
    p = document.add_paragraph()
    p.add_run("Mode de rapport : ").bold = True
    p.add_run(mode_label)

    selected = studio.selected_scenarios()
    document.add_heading("6.1 Scénarios retenus", level=2)
    if selected:
        for sel in selected:
            scenario = SCENARIOS_BY_CODE.get(sel.code)
            title = scenario.libelle if scenario else sel.code
            horizon = f" — {scenario.horizon}" if scenario and scenario.horizon else ""
            document.add_heading(f"{title}{horizon}", level=3)
            if scenario and scenario.description:
                document.add_paragraph(scenario.description)
            if sel.commentaire:
                p = document.add_paragraph()
                p.add_run("Justification OPT'HELIOS : ").bold = True
                p.add_run(_safe_str(sel.commentaire))
            if scenario and scenario.actions_types:
                document.add_paragraph("Actions types associées :")
                _add_bullets(document, list(scenario.actions_types))
    else:
        document.add_paragraph("Aucun scénario stratégique n'a été retenu à ce stade.")

    document.add_heading("6.2 Formulations OPT'HELIOS appliquées", level=2)
    if studio.formulations:
        for idx, applied in enumerate(studio.formulations, start=1):
            template = FORMULATIONS_BY_CODE.get(applied.code)
            title = template.titre if template else applied.code
            section = applied.section or (template.theme if template else "")
            document.add_heading(f"{idx}. {title} — {section}".strip(" —"), level=3)

            constat = applied.constat_personnalise or (template.constat if template else "")
            impact = applied.impact_personnalise or (template.impact if template else "")
            recommandation = applied.recommandation_personnalisee or (
                template.recommandation if template else ""
            )

            if constat:
                p = document.add_paragraph()
                p.add_run("Constat : ").bold = True
                p.add_run(_safe_str(constat))
            if impact:
                p = document.add_paragraph()
                p.add_run("Impact : ").bold = True
                p.add_run(_safe_str(impact))
            if recommandation:
                p = document.add_paragraph()
                p.add_run("Recommandation : ").bold = True
                p.add_run(_safe_str(recommandation))
    else:
        document.add_paragraph("Aucune formulation type OPT'HELIOS n'a été appliquée.")

    if studio.note_strategique:
        document.add_heading("6.3 Note stratégique", level=2)
        document.add_paragraph(_safe_str(studio.note_strategique))


def _add_appendix_metadata(document: Document, metadata: Mapping[str, Any] | None = None) -> None:
    document.add_heading("7. Métadonnées", level=1)

    if not metadata:
        document.add_paragraph("Aucune métadonnée d’audit disponible.")
        return

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    _set_cell_text(table.rows[0].cells[0], "Clé", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Valeur", bold=True)
    _shade_cell(table.rows[0].cells[0], "D9EAF7")
    _shade_cell(table.rows[0].cells[1], "D9EAF7")

    for key, value in metadata.items():
        row = table.add_row().cells
        _set_cell_text(row[0], _safe_str(key))
        _set_cell_text(row[1], _safe_str(value))


def build_docx_report(
    session_state: Any,
    output_path: str | Path,
    *,
    contexte_technique: Mapping[str, Any] | None = None,
    report_title: str = "Rapport d’audit technique solaire thermique",
    site_name: str | None = None,
    reference: str | None = None,
    audit_date: str | None = None,
    include_evidences: bool = True,
    cover_photo_path: str | None = None,
) -> Path:
    payload = build_report_data(session_state, contexte_technique=contexte_technique)

    document = Document()
    _configure_page(document)
    _set_default_font(document)
    _set_document_language(document, "fr-FR")
    _add_header(document, report_title=report_title, reference=reference)
    _add_footer(document)

    _add_cover_page(
        document,
        report_title=report_title,
        site_name=site_name,
        reference=reference,
        audit_date=audit_date,
        cover_photo_path=cover_photo_path,
    )
    document.add_page_break()

    _add_global_assessment(document, payload)
    document.add_page_break()

    _add_executive_summary(document, payload)
    _add_expert_conclusion_section(document, payload)
    document.add_page_break()

    _add_section_summary(document, payload)
    document.add_page_break()

    _add_findings(document, payload, include_evidences=include_evidences)
    document.add_page_break()

    _add_action_plan(document, payload)

    studio = extract_studio_from_session(session_state)
    if studio is not None:
        document.add_page_break()
        _add_studio_sections(document, studio)

    metadata = payload.get("metadata") or {}
    if metadata:
        document.add_page_break()
        _add_appendix_metadata(document, metadata)

    document.add_page_break()
    _add_signature_section(document, auditeur=metadata.get("auditeur"))

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def build_checklist_docx(
    audit: Any,
    output_path: str | Path,
    *,
    report_title: str = "Checklist terrain — Audit solaire thermique",
) -> Path:
    """Checklist condensée à imprimer et remplir à la main PENDANT la visite,
    avant la saisie détaillée dans l'application (contrairement au rapport
    final, qui ne liste que les non-conformités déjà enregistrées, celle-ci
    liste TOUS les points applicables avec des cases vides à cocher)."""
    document = Document()
    _configure_page(document)
    _set_default_font(document)
    _set_document_language(document, "fr-FR")
    _add_footer(document)

    if Path(LOGO_PATH).exists():
        _add_picture_if_exists(document, LOGO_PATH, width_inches=1.4)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report_title)
    run.bold = True
    run.font.size = Pt(16)

    site_name = _safe_str(audit.projet.operation) or _safe_str(audit.projet.adresse.commune)
    if site_name:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(site_name).bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"Date de la visite : ______________    |    Auditeur : ______________________"
    )
    document.add_paragraph("")

    legend = document.add_paragraph()
    legend.add_run("Légende verdict : ").bold = True
    legend.add_run("C = conforme · NC = non conforme · NP = non présent · NV = non vérifiable · SO = sans objet")
    document.add_paragraph("")

    grouped = group_controls_by_section(audit)
    for section in sorted(grouped.keys()):
        items = grouped[section]

        document.add_heading(section, level=1)

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = ["ID", "Point de contrôle", "Verdict", "Observations"]
        for cell, label in zip(table.rows[0].cells, headers):
            _set_cell_text(cell, label, bold=True)
            _shade_cell(cell, "D9EAF7")

        for item in items:
            row = table.add_row().cells
            _set_cell_text(row[0], item.controle_id)
            _set_cell_text(row[1], item.libelle)
            _set_cell_text(row[2], "")
            _set_cell_text(row[3], "")

        document.add_paragraph("")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
